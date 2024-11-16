import os
from docx import Document
import re

# Function to combine text files based on name order
def combine_text_files(folder_path, output_file):
    with open(output_file, "w", encoding="utf-8") as outfile:
        file_names = [f for f in os.listdir(folder_path) if f.endswith(".txt")]
        file_names_sorted = sorted(file_names, key=lambda x: [int(num) for num in re.findall(r'\d+', x)])
        for file_name in file_names_sorted:
            with open(os.path.join(folder_path, file_name), "r", encoding="utf-8") as infile:
                outfile.write("\n")
                outfile.write(infile.read())
                outfile.write("\n")



doc = Document()

# Function to combine text files based on name order
def combine_text_files_to_word(folder_path, output_file):
    with open(output_file, "w", encoding="utf-8") as outfile:
        file_names = [f for f in os.listdir(folder_path) if f.endswith(".txt")]
        file_names_sorted = sorted(file_names, key=lambda x: [int(num) for num in re.findall(r'\d+', x)])
        for file_name in file_names_sorted:
            with open(os.path.join(folder_path, file_name), "r", encoding="utf-8") as infile:
                doc.add_paragraph(infile.read())
                doc.add_paragraph("\n")

        doc.save(output_file)



# combine_text_files("output/bookname", "output/bookname/bookname.txt")
# combine_text_files_to_word("output/bookname", "output/bookname/bookname.docx")
# print("Text files combined successfully!")
