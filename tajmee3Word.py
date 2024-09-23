import os
from docx import Document
import re
# Path to the folder containing the text files
folder_path = "output/المقامر رواية لـ دوستويفسكي_PDF"

# Output file where the combined content will be saved
output_file = f"{folder_path}/!combined_files.docx"

doc = Document()

# Function to combine text files based on name order
def combine_text_files(folder_path, output_file):
    with open(output_file, "w", encoding="utf-8") as outfile:
        file_names = [f for f in os.listdir(folder_path) if f.endswith(".txt")]
        file_names_sorted = sorted(file_names, key=lambda x: [int(num) for num in re.findall(r'\d+', x)])
        for file_name in file_names_sorted:
            with open(os.path.join(folder_path, file_name), "r", encoding="utf-8") as infile:
                Page_number = re.findall(r'\d+', file_name)
                Page_number = ''.join(Page_number)
                doc.add_paragraph(f"Page#{Page_number}\n")
                doc.add_paragraph(infile.read())
                doc.add_paragraph("\n")

        doc.save(output_file)

# Example usage
combine_text_files(folder_path, output_file)
print("Text files combined successfully!")
